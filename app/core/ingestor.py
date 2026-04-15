"""
Ingestor: parses documents into structured data for indexing and retrieval.

Responsibilities:
  - load_caiq_questions(): reads the CAIQ Excel file (legacy flow).
  - load_customer_docs(): reads customer PDFs/XLSX for text (legacy flow).
  - parse_document(): universal parser for docx/pdf/xlsx files.
  - chunk_text(): word-based overlapping chunker (~400 tokens / 100 overlap).
  - load_sop_file(): parse + chunk a single SOP file with a capability tag.
  - load_sow_file(): parse + chunk a single SOW file.
  - load_questions_xlsx(): parse a questions Excel (category + question cols).
"""

import re
from pathlib import Path
from typing import List, Dict

import openpyxl
import fitz  # PyMuPDF


# --- Constants ---

# Valid CAIQ question IDs follow the pattern: DOMAIN-##.## (e.g. IAM-01.1, A&A-03.2)
VALID_ID_PATTERN = re.compile(r"^[A-Z&]+-\d+\.\d+$")


# ---------------------------------------------------------------------------
# CAIQ XLSX Parser
# ---------------------------------------------------------------------------

def load_caiq_questions(xlsx_path: str) -> List[Dict]:
    """
    Parse the CAIQ XLSX file and return a list of question dicts.

    Reads the 'CAIQv4.0.3' sheet, skips header/footer rows using
    VALID_ID_PATTERN, and extracts the question ID, domain (derived
    from the ID prefix), and cleaned question text.

    Args:
        xlsx_path: absolute or relative path to the CAIQ .xlsx file.

    Returns:
        List of dicts, each with keys:
            question_id  — e.g. "IAM-01.1"
            domain       — e.g. "IAM"  (prefix before the first dash)
            question_text — cleaned question string
            source       — stem of the source file name
    """
    # Open workbook in read-only mode for memory efficiency
    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    ws = wb["CAIQv4.0.3"]

    questions = []
    for row in ws.iter_rows(values_only=True):
        question_id = row[0]
        question_text = row[1]

        # Skip blank rows and non-question rows (headers, footers, totals)
        if not question_id or not question_text:
            continue
        if not VALID_ID_PATTERN.match(str(question_id).strip()): #check by removing whitespace and converting to string in case of numeric IDs
            continue

        # Clean up whitespace and inline newlines from the cell value
        question_id = str(question_id).strip()
        domain = question_id.split("-")[0]          # e.g. "IAM" from "IAM-01.1"
        question_text = str(question_text).strip().replace("\n", " ")

        questions.append({
            "question_id": question_id,
            "domain": domain,
            "question_text": question_text,
            "source": Path(xlsx_path).stem,         # filename without extension
        })

    wb.close()
    return questions


# ---------------------------------------------------------------------------
# Customer PDF Parser
# ---------------------------------------------------------------------------

def load_customer_docs(customer_dir: str) -> str:
    """
    Parse all PDFs and XLSX files in a customer directory and return their combined text.

    Each PDF is opened with PyMuPDF; each XLSX is read with openpyxl (all
    non-empty cell values concatenated row by row). Results are concatenated
    with a filename header so downstream code can identify which doc each
    chunk came from.

    Args:
        customer_dir: path to folder containing the customer's PDF/XLSX files.

    Returns:
        Single string with all documents concatenated, separated by headers.

    Raises:
        FileNotFoundError: if no supported files are found in the directory.
    """
    customer_path = Path(customer_dir)

    pdf_files = sorted(customer_path.glob("*.pdf"))
    xlsx_files = sorted(customer_path.glob("*.xlsx"))
    all_files = pdf_files + xlsx_files

    if not all_files:
        raise FileNotFoundError(f"No PDF or XLSX files found in {customer_dir}")

    combined_text = []

    for pdf_file in pdf_files:
        doc = fitz.open(str(pdf_file))
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        combined_text.append(f"=== {pdf_file.name} ===\n{text.strip()}")

    for xlsx_file in xlsx_files:
        wb = openpyxl.load_workbook(xlsx_file, read_only=True, data_only=True)
        rows_text = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if row_values:
                    rows_text.append(" | ".join(row_values))
        wb.close()
        combined_text.append(f"=== {xlsx_file.name} ===\n" + "\n".join(rows_text))

    return "\n\n".join(combined_text)


# ---------------------------------------------------------------------------
# Universal document parser
# ---------------------------------------------------------------------------

def parse_document(file_path: str) -> str:
    """
    Parse a document file (docx, pdf, xlsx) and return its plain text.

    Supports:
      - .docx — Microsoft Word via python-docx
      - .pdf  — via PyMuPDF (fitz)
      - .xlsx — via openpyxl (all cell values joined)

    Args:
        file_path: path to the document file.

    Returns:
        Extracted text as a single string.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        from docx import Document  # lazy import to avoid hard dep at module load
        doc = Document(str(path))
        lines = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table cell text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)

    elif suffix == ".pdf":
        doc = fitz.open(str(path))
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text

    elif suffix == ".xlsx":
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        rows_text = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_values = [
                    str(cell).strip() for cell in row
                    if cell is not None and str(cell).strip()
                ]
                if row_values:
                    rows_text.append(" | ".join(row_values))
        wb.close()
        return "\n".join(rows_text)

    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# ---------------------------------------------------------------------------
# Text chunker (word-based, ~400 tokens / 100 overlap)
# ---------------------------------------------------------------------------

def chunk_text(text: str, chunk_size: int = 300, overlap: int = 75) -> List[str]:
    """
    Split text into overlapping word-based chunks.

    Approximates token counts (1 token ≈ ~1.3 words for English):
      - chunk_size=300 words ≈ 400 tokens
      - overlap=75 words ≈ 100 tokens

    Args:
        text:       input text to chunk.
        chunk_size: words per chunk (default 300 ≈ 400 tokens).
        overlap:    word overlap between consecutive chunks (default 75).

    Returns:
        List of non-empty chunk strings.
    """
    words = text.split()
    if not words:
        return []

    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk)
        if end >= len(words):
            break
        start += chunk_size - overlap

    return chunks


# ---------------------------------------------------------------------------
# SOP file loader (parse + chunk with capability tag)
# ---------------------------------------------------------------------------

def load_sop_file(file_path: str, capability: str) -> List[Dict]:
    """
    Parse a SOP file and return overlapping chunks with metadata.

    Args:
        file_path:  path to SOP file (.docx, .pdf, or .xlsx).
        capability: user-supplied capability label for this SOP (e.g. "IAM").

    Returns:
        List of dicts, each with:
            chunk_id     — "<stem>_<index>"
            filename     — original filename
            capability   — capability label
            chunk_text   — text of this chunk
            chunk_index  — 0-based chunk position
    """
    path = Path(file_path)
    text = parse_document(file_path)
    chunks = chunk_text(text)

    return [
        {
            "chunk_id": f"{path.stem}_{i}",
            "filename": path.name,
            "capability": capability.strip(),
            "chunk_text": chunk,
            "chunk_index": i,
        }
        for i, chunk in enumerate(chunks)
    ]


# ---------------------------------------------------------------------------
# SOW file loader (parse + chunk, no capability tag)
# ---------------------------------------------------------------------------

def load_sow_file(file_path: str) -> List[Dict]:
    """
    Parse a SOW file and return overlapping chunks.

    Args:
        file_path: path to SOW file (.docx, .pdf, or .xlsx).

    Returns:
        List of dicts, each with:
            chunk_id    — "<stem>_<index>"
            filename    — original filename
            chunk_text  — text of this chunk
            chunk_index — 0-based chunk position
    """
    path = Path(file_path)
    text = parse_document(file_path)
    chunks = chunk_text(text)

    return [
        {
            "chunk_id": f"{path.stem}_{i}",
            "filename": path.name,
            "chunk_text": chunk,
            "chunk_index": i,
        }
        for i, chunk in enumerate(chunks)
    ]


# ---------------------------------------------------------------------------
# Questions Excel parser
# ---------------------------------------------------------------------------

def load_questions_xlsx(xlsx_path: str) -> List[Dict]:
    """
    Parse a questions Excel file with 'category' and 'question' columns.

    Automatically detects the header row by looking for cells containing
    'category' and 'question' (case-insensitive).  Defaults to columns
    0 and 1 if no header is found.

    Args:
        xlsx_path: path to the Excel file.

    Returns:
        List of dicts, each with:
            question_id   — auto-generated (e.g. "IAM_001")
            category      — category / domain string
            question_text — question text
    """
    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    # Detect header row and column positions
    cat_col, q_col = 0, 1
    header_row_idx = 0

    for idx, row in enumerate(rows):
        if not row:
            continue
        headers = [str(cell).strip().lower() if cell is not None else "" for cell in row]
        cat_found = next((i for i, h in enumerate(headers) if "category" in h), None)
        q_found = next(
            (i for i, h in enumerate(headers) if "question" in h and i != cat_found),
            None,
        )
        if cat_found is not None and q_found is not None:
            cat_col = cat_found
            q_col = q_found
            header_row_idx = idx
            break

    questions: List[Dict] = []
    question_counter: Dict[str, int] = {}

    for row in rows[header_row_idx + 1 :]:
        if not row or len(row) <= max(cat_col, q_col):
            continue

        category = row[cat_col]
        question_text = row[q_col]

        if not category or not question_text:
            continue

        category = str(category).strip()
        question_text = str(question_text).strip()

        if not category or not question_text:
            continue

        # Build a short stable key from the category name
        cat_key = re.sub(r"[^A-Z0-9]", "_", category[:15].upper()).strip("_")
        question_counter[cat_key] = question_counter.get(cat_key, 0) + 1
        question_id = f"{cat_key}_{question_counter[cat_key]:03d}"

        questions.append(
            {
                "question_id": question_id,
                "category": category,
                "question_text": question_text,
            }
        )

    return questions
