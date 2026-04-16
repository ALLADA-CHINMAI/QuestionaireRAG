"""
Convert SOW text files to Word documents.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_docx_from_text(text_file_path: str, output_docx_path: str):
    """Convert a text SOW file to a formatted Word document."""
    
    # Read the text content
    with open(text_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content into lines
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - add space
            doc.add_paragraph()
        elif line.startswith('Statement of Work No.'):
            # Title
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
        elif line.startswith('Customer1'):
            # Subtitle
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
        elif line.isupper() and len(line) < 60:
            # Section headers (all caps)
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(12)
        elif line[0].isdigit() and '. ' in line[:5]:
            # Numbered list items
            p = doc.add_paragraph(line, style='List Number')
        elif line.startswith('- '):
            # Bullet points
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif '|' in line and '-' * 5 in line:
            # Table separator line - skip
            continue
        elif '|' in line:
            # Table row - handle separately
            # For simplicity, just add as paragraph
            doc.add_paragraph(line)
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save the document
    doc.save(output_docx_path)
    print(f"Created: {output_docx_path}")


if __name__ == "__main__":
    data_dir = Path(__file__).parent.parent / "data"
    
    # Convert SOW No. 26
    create_docx_from_text(
        str(data_dir / "Customer1 SOW No. 26 - Security Compliance Assessment_Draft.txt"),
        str(data_dir / "Customer1 SOW No. 26 - Security Compliance Assessment_Draft.docx")
    )
    
    # Convert SOW No. 27
    create_docx_from_text(
        str(data_dir / "Customer1 SOW No. 27 - Data Engineering Operations_Draft.txt"),
        str(data_dir / "Customer1 SOW No. 27 - Data Engineering Operations_Draft.docx")
    )
    
    print("\n✓ Both SOW documents converted to .docx format")
